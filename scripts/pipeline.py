#!/usr/bin/env python3
"""
Amazon Review Pipeline - 全链路自动化
1搜索ASIN → 2产品数据 → 3评论抓取 → 4AI翻译 → 5分析报告

与 cdp_scraper.py / cdp_scraper_reviews.py 保持一致的 CDP 写法:
  - get_cdp_ws_url() 用 /json/list 获取页面级 WS URL
  - cdp_send() 返回 result.value(returnByValue=True)
  - 产品/评论提取均用 JS DOM 而非 innerHTML 正则
"""

import argparse, asyncio, json, os, random, re, subprocess, sys, time, uuid
from pathlib import Path
import requests, websockets

# ─── Region / Domain / Language ─────────────────────────────────────────────
AMAZON_DOMAINS = {
    "us": "www.amazon.com", "uk": "www.amazon.co.uk", "de": "www.amazon.de",
    "jp": "www.amazon.co.jp", "fr": "www.amazon.fr", "it": "www.amazon.it",
    "es": "www.amazon.es", "ca": "www.amazon.ca", "in": "www.amazon.in",
    "au": "www.amazon.com.au", "mx": "www.amazon.com.mx", "br": "www.amazon.com.br",
    "nl": "www.amazon.nl",
}

REGION_NAMES = {
    "us": "美国", "uk": "英国", "de": "德国", "jp": "日本", "fr": "法国",
    "it": "意大利", "es": "西班牙", "ca": "加拿大", "in": "印度",
    "au": "澳大利亚", "mx": "墨西哥", "br": "巴西", "nl": "荷兰",
}

SOURCE_LANGUAGES = {
    "us": "English", "uk": "English", "ca": "English", "au": "English", "in": "English",
    "de": "German", "jp": "Japanese", "fr": "French", "it": "Italian",
    "es": "Spanish", "mx": "Spanish", "br": "Portuguese",
    "nl": "Dutch",
}

# ─── Config ──────────────────────────────────────────────────────────────────
DEFAULT_PORT = int(os.environ.get("CDP_PORT", "18800"))
FEISHU_WEBHOOK = os.environ.get("FEISHU_WEBHOOK", "")
MINIMAX_BASE_URL = os.environ.get("MINIMAX_BASE_URL", "https://api.minimaxi.com/v1")
DEFAULT_ATTACHMENT_ROOT = os.environ.get(
    "AMAZON_REVIEW_PIPELINE_ATTACHMENTS",
    str(Path.home() / "Documents" / "Obsidian" / "Amazon选品" / "99_附件"),
)
AI_REQUEST_TIMEOUT = int(os.environ.get("AMAZON_REVIEW_PIPELINE_AI_TIMEOUT", "300"))
AI_MAX_RETRIES = int(os.environ.get("AMAZON_REVIEW_PIPELINE_AI_RETRIES", "3"))

def _load_minimax_key() -> str:
    """加载 MiniMax API Key：优先环境变量，其次 openclaw.json 配置"""
    key = os.environ.get("MINIMAX_API_KEY", "")
    if key:
        return key
    # 从 openclaw.json 的 env 配置中读取
    for cfg_path in [
        os.path.expanduser("~/.openclaw/openclaw.json"),
        os.path.expanduser("~/.openclaw/config.json"),
    ]:
        if os.path.exists(cfg_path):
            try:
                with open(cfg_path, "r", encoding="utf-8") as f:
                    cfg = json.load(f)
                # openclaw.json 里 key 可能在 env 对象或顶层
                key = cfg.get("env", {}).get("MINIMAX_API_KEY", "") if isinstance(cfg.get("env"), dict) else ""
                if not key:
                    key = cfg.get("MINIMAX_API_KEY", "")
                if key:
                    return key
            except Exception:
                continue
    return ""

API_KEY = _load_minimax_key()
TRANSLATE_TIMEOUT = int(os.environ.get("AMAZON_REVIEW_PIPELINE_TRANSLATE_TIMEOUT", "90"))
RETRY_MAX = 3
BATCH_SIZE = 3
BATCH_DELAY = 0.5

# ─── 中文关键词自动翻译 ───────────────────────────────────────────────────
KEYWORD_TRANSLATIONS = {
    "抱枕": "body pillow",
    "儿童棉袜": "kids cotton socks",
    "儿童袜子": "kids socks",
    "儿童运动袜": "kids athletic socks",
    "瑜伽裤": "yoga pants",
    "运动短裤": "athletic shorts women",
    "运动内衣": "sports bra women",
    "跑步短裤": "running shorts women",
    "健身服": "fitness wear women",
    "羽毛球拍": "badminton racket",
    "网球拍": "tennis racket",
    "游泳镜": "swimming goggles",
    "保温杯": "insulated water bottle",
    "水杯": "water bottle",
    "杯子": "water bottle",
    "塑料水杯": "plastic water bottle",
    "无线充电器": "wireless charger",
    "蓝牙耳机": "bluetooth earphones",
    "充电线": "charging cable",
    "纸质笔记本": "journal notebook",
}

def auto_translate_keyword(kw: str) -> tuple[str, str]:
    """检测中文并翻译为英文,返回 (原始关键词, 搜索用关键词)"""
    has_chinese = any("\u4e00" <= c <= "\u9fff" for c in kw)
    if not has_chinese:
        return kw, kw
    en = KEYWORD_TRANSLATIONS.get(kw)
    if en:
        print(f"   [WEB] 检测到中文关键词 -> 自动翻译为: {en}")
        return kw, en
    print(f"   [!] 未找到 '{kw}' 的翻译,请使用英文关键词")
    return kw, kw

# ─── 品类显示(仅 emoji + 名称)─────────────────────────────────────────────
CATEGORY_EMOJI = {"充电线": "🔌", "塑料水杯": "🥤", "抱枕": "🛏️", "纸质笔记本": "📓"}

def get_category_display(keyword: str) -> tuple:
    for cat, emoji in CATEGORY_EMOJI.items():
        if cat in keyword or keyword in cat:
            return emoji, cat
    return "📦", "商品"

CDP_BROWSER = None   # 全局当前 Chrome 进程
CDP_TAB_IDS = {}
_WARMED_DOMAINS = set()

# ─── Chrome / CDP ────────────────────────────────────────────────────────────
# 使用已登录的 Chrome 调试端口,直接导航到各 ASIN 评论页


def open_new_tab(port: int, url: str) -> str:
    """通过 PUT /json/new 创建新标签页并直接导航到指定 URL，返回 WS URL"""
    for attempt in range(5):
        try:
            resp = requests.put(
                f"http://localhost:{port}/json/new?{url}", timeout=10)
            if resp.status_code == 200:
                tab = resp.json()
                ws = tab.get("webSocketDebuggerUrl", "")
                if ws:
                    CDP_TAB_IDS[ws] = tab.get("id", "")
                    return ws
            elif resp.status_code == 500 and attempt < 4:
                # Chrome 偶发 500，重试
                time.sleep(1.5)
                continue
        except Exception:
            if attempt < 4:
                time.sleep(1.5)
                continue
        break
    raise RuntimeError(f"无法创建新标签页 on port {port}")


def close_tab(port: int, ws_url: str):
    tab_id = CDP_TAB_IDS.pop(ws_url, "")
    if not tab_id:
        return
    try:
        requests.get(f"http://localhost:{port}/json/close/{tab_id}", timeout=3)
    except Exception:
        pass


async def warmup_domain(domain: str):
    """进入搜索/榜单前先访问首页，建立正常会话和 referer，避免被当成异常访问。"""
    if domain in _WARMED_DOMAINS:
        return
    port = DEFAULT_PORT
    print(f"   🔥 预热 {domain} 首页…")
    ws = open_new_tab(port, f"https://{domain}/")
    try:
        await asyncio.sleep(random.uniform(5, 8))
        await cdp_send(ws, "Runtime.evaluate", {"expression": "window.scrollTo(0, Math.max(400, document.body.scrollHeight/3))"})
        await asyncio.sleep(random.uniform(1.5, 3.5))
    finally:
        close_tab(port, ws)
    _WARMED_DOMAINS.add(domain)


def get_cdp_ws_url(port: int, url_contains: str = "") -> str:
    """从 /json/list 获取当前活跃标签页的页面级 WebSocket URL
    url_contains: 可选，优先匹配 URL 包含此字符串的标签页"""
    for attempt in range(20):
        try:
            info = json.loads(
                requests.get(f"http://localhost:{port}/json/list", timeout=3).text)
            # 如果指定了 url_contains，优先匹配
            if url_contains:
                for t in info:
                    url = t.get("url", "")
                    if t.get("type") == "page" and url_contains in url:
                        return t["webSocketDebuggerUrl"]
            # fallback: 第一个非广告页面
            for t in info:
                url = t.get("url", "")
                if t.get("type") == "page" and "amazon-adsystem" not in url and url:
                    return t["webSocketDebuggerUrl"]
        except Exception:
            pass
        # 第一次尝试失败后，如果超过 5 秒还没找到，尝试打开新标签页
        if attempt == 3:
            try:
                resp = requests.put(
                    f"http://localhost:{port}/json/new?about:blank", timeout=5)
                if resp.status_code == 200:
                    new_tab = resp.json()
                    if new_tab.get("webSocketDebuggerUrl"):
                        return new_tab["webSocketDebuggerUrl"]
            except Exception:
                pass
        time.sleep(1)
    raise RuntimeError(f"无法连接 Chrome CDP on port {port}")


# ─── CDP 底层 ────────────────────────────────────────────────────────────────
async def cdp_send(ws_url: str, method: str, params: dict = None, timeout: int = 30):
    """
    发送 CDP 命令,返回 result.value(与 cdp_scraper.py 行为一致)
    returnByValue=True 让 JS 返回实际值而非引用
    """
    params = dict(params or {})
    params["returnByValue"] = True
    msg_id = random.randint(1, 999999)

    last_err = None
    for wsa in range(4):  # WS 连接重试 4 次
        try:
            async with websockets.connect(ws_url, max_size=50 * 1024 * 1024) as conn:
                await conn.send(json.dumps({"id": msg_id, "method": method, "params": params}))
                for _ in range(timeout * 2):
                    try:
                        raw = await asyncio.wait_for(conn.recv(), timeout=0.5)
                        msg = json.loads(raw)
                        if msg.get("id") == msg_id:
                            r = msg.get("result", {})
                            # 兼容 result.result.value 或 result.value
                            if "result" in r:
                                return r["result"].get("value", "")
                            return r.get("value", "")
                    except asyncio.TimeoutError:
                        continue
        except Exception as e:
            last_err = e
            if wsa < 3:
                await asyncio.sleep(2 ** wsa)  # 指数退避 1s, 2s, 4s
                continue
        break

    return ""


# ─── 翻译 ────────────────────────────────────────────────────────────────────
def translate_batch_m2(reviews: list[dict], batch_size: int = 5, source_lang: str = "English") -> list[dict]:
    """MiniMax M2.7 批量翻译:多语言→中,字段 title_zh / body_zh"""
    if not reviews:
        return reviews
    results = []
    headers = {"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"}

    for i in range(0, len(reviews), batch_size):
        batch = reviews[i:i + batch_size]

        # ── 重试逻辑：对每条评论独立重试翻译 ─────────────────────
        for retry_round in range(3):
            untranslated = [r for r in batch if not r.get("body_zh") or r["body_zh"] == r.get("body", "")]
            if not untranslated:
                break

            lines = [f"[{j}] Title: {r.get('title','').replace(chr(10),' ').strip()}\n"
                     f"[{j}] Body: {r.get('body','').replace(chr(10),' ').strip()}"
                     for j, r in enumerate(batch)]

            payload = {
                "model": "MiniMax-M2.7",
                "messages": [{"role": "user", "content":
                    f"You are a professional translator. Translate Amazon reviews from {source_lang} "
                    "to Simplified Chinese. Output ONLY valid JSON array.\n"
                    'Format: [{"idx":N,"title_zh":"...","body_zh":"..."},...]\n'
                    "Reviews:\n" + "\n".join(lines)}],
                "temperature": 0.3, "max_tokens": 2000,
                "do_sample": True, "thinking": False,
            }

            ok = False
            try:
                resp = requests.post(
                    f"{MINIMAX_BASE_URL}/chat/completions",
                    headers=headers, json=payload, timeout=TRANSLATE_TIMEOUT, verify=False)
                if resp.status_code == 200:
                    try:
                        raw = resp.json()["choices"][0]["message"]["content"].strip()
                        # MiniMax M2.7 推理模型会在 content 中输出 <think...>...</think...> 标签包裹思考过程
                        # 需要先提取思考标签之后的有效内容
                        think_end = re.search(r'</think\s*>', raw, re.IGNORECASE | re.DOTALL)
                        if think_end:
                            raw = raw[think_end.end():].strip()
                        # 去除残留的 HTML 标签（如 <工具调用> 等）
                        content2 = re.sub(r'<[^>]+>', '', raw).strip()
                        # 去除代码围栏
                        fence_match = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', content2, re.DOTALL)
                        if fence_match:
                            content2 = fence_match.group(1).strip()
                        translations = json.loads(content2)
                        trans_map = {str(t["idx"]): t for t in translations}
                        for idx, r in enumerate(batch):
                            t = trans_map.get(str(idx)) or {}
                            body_zh = t.get("body_zh", "").strip()
                            title_zh = t.get("title_zh", "").strip()
                            if body_zh and body_zh != r.get("body", ""):
                                r["body_zh"] = body_zh
                                r["title_zh"] = title_zh or r.get("title", "")
                        ok = True
                    except Exception:
                        pass
            except Exception:
                pass

            if not ok:
                time.sleep(1)

        # 兜底：仍未翻译的用原文
        for r in batch:
            r["body_zh"]  = r.get("body_zh",  "") or r.get("body",  "")
            r["title_zh"] = r.get("title_zh", "") or r.get("title", "")
            results.append(r)

        print(f"    ✅ 翻译 {len(batch)} 条")
        time.sleep(0.5)

    return results


def parse_price_value(price: str):
    text = (price or "").replace(",", "")
    match = re.search(r"\d+(?:\.\d+)?", text)
    if not match:
        return None
    try:
        return float(match.group())
    except ValueError:
        return None


def make_run_output_dir(run_id: str):
    path = Path(DEFAULT_ATTACHMENT_ROOT).expanduser() / run_id
    path.mkdir(parents=True, exist_ok=True)
    return path


LOCALIZED_KEYWORD_HINTS = {
    "us": {
        "water bottle": ["reusable water bottle", "insulated water bottle"],
        "bottle": ["reusable water bottle", "insulated water bottle"],
    },
    "uk": {
        "water bottle": ["reusable water bottle", "insulated water bottle"],
        "bottle": ["reusable water bottle", "insulated water bottle"],
    },
    "jp": {
        "water bottle": ["水筒", "ウォーターボトル"],
        "bottle": ["水筒", "ウォーターボトル"],
    },
    "de": {
        "water bottle": ["Trinkflasche", "Isolierflasche"],
        "bottle": ["Trinkflasche", "Isolierflasche"],
    },
    "fr": {
        "water bottle": ["gourde", "bouteille d'eau réutilisable"],
        "bottle": ["gourde", "bouteille réutilisable"],
    },
    "it": {
        "water bottle": ["borraccia", "bottiglia acqua riutilizzabile"],
        "bottle": ["borraccia"],
    },
    "es": {
        "water bottle": ["botella de agua reutilizable", "termo agua"],
        "bottle": ["botella de agua reutilizable"],
    },
    "ca": {
        "water bottle": ["reusable water bottle", "insulated water bottle"],
        "bottle": ["reusable water bottle", "insulated water bottle"],
    },
    "in": {
        "water bottle": ["reusable water bottle", "insulated water bottle"],
        "bottle": ["reusable water bottle", "insulated water bottle"],
    },
    "au": {
        "water bottle": ["reusable water bottle", "insulated water bottle"],
        "bottle": ["reusable water bottle", "insulated water bottle"],
    },
    "mx": {
        "water bottle": ["botella de agua reutilizable", "termo para agua"],
        "bottle": ["botella de agua reutilizable"],
    },
    "br": {
        "water bottle": ["garrafa de água reutilizável", "squeeze garrafa"],
        "bottle": ["garrafa de água reutilizável"],
    },
    "nl": {
        "water bottle": ["drinkfles", "waterfles herbruikbaar"],
        "bottle": ["drinkfles"],
    },
}


SEARCH_DRIFT_RULES = {
    "us": {
        "water bottle": {
            "bad_terms": ["24 pack", "case of water", "spring water", "purified water", "mineral water", "bottled water"],
            "suggestions": ["reusable water bottle", "insulated water bottle"],
            "message": "当前搜索结果可能偏向瓶装饮用水，而不是可复用水杯/水壶。",
        },
    },
    "uk": {
        "water bottle": {
            "bad_terms": ["24 pack", "case of water", "spring water", "mineral water", "bottled water"],
            "suggestions": ["reusable water bottle", "insulated water bottle"],
            "message": "当前英国站搜索结果可能偏向瓶装饮用水，而不是可复用水杯/水壶。",
        },
    },
    "jp": {
        "water bottle": {
            "bad_terms": ["天然水", "ミネラルウォーター", "ラベルレス", "PET", "24本", "水 ミネラル"],
            "suggestions": ["水筒", "ウォーターボトル"],
            "message": "当前日本站搜索结果可能偏向瓶装水/矿泉水，而不是水杯/水壶。",
        },
    },
    "de": {
        "water bottle": {
            "bad_terms": ["mineralwasser", "wasser still", "wasser sprudel", "pfandflasche", "24 x"],
            "suggestions": ["Trinkflasche", "Isolierflasche"],
            "message": "当前德国站搜索结果可能偏向瓶装饮用水，而不是可复用水杯/水壶。",
        },
    },
    "fr": {
        "water bottle": {
            "bad_terms": ["eau minérale", "pack d'eau", "bouteilles d'eau", "eau de source", "24 x"],
            "suggestions": ["gourde", "bouteille d'eau réutilisable"],
            "message": "当前法国站搜索结果可能偏向瓶装饮用水，而不是可复用水杯/水壶。",
        },
    },
    "it": {
        "water bottle": {
            "bad_terms": ["acqua minerale", "confezione", "bottiglie acqua", "24 x", "naturale"],
            "suggestions": ["borraccia", "bottiglia acqua riutilizzabile"],
            "message": "当前意大利站搜索结果可能偏向瓶装饮用水，而不是可复用水杯/水壶。",
        },
    },
    "es": {
        "water bottle": {
            "bad_terms": ["agua mineral", "pack de agua", "botellas de agua", "24 x", "garrafas"],
            "suggestions": ["botella de agua reutilizable", "termo agua"],
            "message": "当前西班牙站搜索结果可能偏向瓶装饮用水，而不是可复用水杯/水壶。",
        },
    },
    "ca": {
        "water bottle": {
            "bad_terms": ["24 pack", "case of water", "spring water", "mineral water", "bottled water"],
            "suggestions": ["reusable water bottle", "insulated water bottle"],
            "message": "当前加拿大站搜索结果可能偏向瓶装饮用水，而不是可复用水杯/水壶。",
        },
    },
    "in": {
        "water bottle": {
            "bad_terms": ["mineral water", "packaged drinking water", "bottled water", "24 pack", "case of water"],
            "suggestions": ["reusable water bottle", "insulated water bottle"],
            "message": "当前印度站搜索结果可能偏向瓶装饮用水，而不是可复用水杯/水壶。",
        },
    },
    "au": {
        "water bottle": {
            "bad_terms": ["24 pack", "case of water", "spring water", "mineral water", "bottled water"],
            "suggestions": ["reusable water bottle", "insulated water bottle"],
            "message": "当前澳大利亚站搜索结果可能偏向瓶装饮用水，而不是可复用水杯/水壶。",
        },
    },
    "mx": {
        "water bottle": {
            "bad_terms": ["agua mineral", "paquete de agua", "botellas de agua", "garrafón", "24 x"],
            "suggestions": ["botella de agua reutilizable", "termo para agua"],
            "message": "当前墨西哥站搜索结果可能偏向瓶装饮用水，而不是可复用水杯/水壶。",
        },
    },
    "br": {
        "water bottle": {
            "bad_terms": ["água mineral", "agua mineral", "fardo", "garrafas de água", "24 x"],
            "suggestions": ["garrafa de água reutilizável", "squeeze garrafa"],
            "message": "当前巴西站搜索结果可能偏向瓶装饮用水，而不是可复用水杯/水壶。",
        },
    },
    "nl": {
        "water bottle": {
            "bad_terms": ["mineraalwater", "bronwater", "flessen water", "24 x"],
            "suggestions": ["drinkfles", "waterfles herbruikbaar"],
            "message": "当前荷兰站搜索结果可能偏向瓶装饮用水，而不是可复用水杯/水壶。",
        },
    },
}


def print_localized_keyword_hint(keyword: str, region: str):
    hints = LOCALIZED_KEYWORD_HINTS.get(region, {}).get(keyword.lower())
    if hints:
        print(f"  💡 关键词本地化建议: {REGION_NAMES.get(region, region)}站可优先尝试 " + " / ".join(hints))


def diagnose_search_drift(products: list[dict], keyword: str, region: str):
    rule = SEARCH_DRIFT_RULES.get(region, {}).get(keyword.lower())
    if not rule or not products:
        return
    titles = [p.get("title", "") for p in products]
    hits = []
    for title in titles:
        if any(term.lower() in title.lower() for term in rule["bad_terms"]):
            hits.append(title)
    ratio = len(hits) / max(len(titles), 1)
    if ratio >= 0.4:
        print("  ⚠️ 关键词质量提示: " + rule["message"])
        print("  ✅ 建议改用: " + " / ".join(rule["suggestions"]))


# ─── AI 语义分析 ────────────────────────────────────────────────────────────
def fallback_review_insights(reviews: list[dict]) -> dict:
    def parse_rating_value(value):
        match = re.search(r"\d+(?:\.\d+)?", str(value or ""))
        return float(match.group(0)) if match else None

    pain_rules = {
        "漏水/密封问题": ["漏", "leak", "leaking", "leaked"],
        "生锈/异味": ["锈", "rust", "smell", "odor", "味"],
        "破损/瑕疵": ["broken", "damaged", "dent", "scratch", "刮", "缺口", "损坏"],
        "配件缺失": ["missing", "缺", "没有", "盖子", "lid"],
        "颜色/外观不符": ["color", "颜色", "picture", "图片", "disappointed", "失望"],
        "尺寸/重量问题": ["heavy", "bulky", "too big", "too small", "重", "太大", "太小"],
    }
    pro_rules = {
        "保冷效果好": ["cold", "ice", "冰", "保冷"],
        "不漏水": ["no leaks", "hasn't leaked", "leak proof", "不会漏", "不漏"],
        "结实耐用": ["durable", "strong", "rugged", "坚固", "耐用"],
        "便携/适合出行": ["travel", "gym", "hike", "portable", "旅行", "健身", "徒步", "携带"],
        "配件/盖子实用": ["lid", "straw", "accessories", "盖", "吸管", "配件"],
        "容量/尺寸合适": ["size", "cup holder", "尺寸", "杯架"],
    }
    scenario_rules = {
        "运动健身": ["gym", "健身"],
        "徒步露营": ["hike", "camping", "徒步", "露营"],
        "通勤上班": ["work", "上班"],
        "旅行出行": ["travel", "旅行"],
        "日常饮水": ["daily", "everywhere", "每天", "饮水"],
        "车载使用": ["car cup holder", "cup holder", "车载", "杯架"],
    }
    pain_counts = {}
    pro_counts = {}
    scenario_counts = {}
    best_review_index = 0
    best_len = 0
    for idx, review in enumerate(reviews or []):
        text = ((review.get("title_zh") or review.get("title") or "") + " " + (review.get("body_zh") or review.get("body") or "")).lower()
        if len(text) > best_len:
            best_len = len(text)
            best_review_index = idx
        rating_number = parse_rating_value(review.get("rating"))
        for key, words in pain_rules.items():
            if any(word.lower() in text for word in words) or (rating_number is not None and rating_number <= 3):
                pain_counts[key] = pain_counts.get(key, 0) + 1
        for key, words in pro_rules.items():
            if any(word.lower() in text for word in words):
                pro_counts[key] = pro_counts.get(key, 0) + 1
        for key, words in scenario_rules.items():
            if any(word.lower() in text for word in words):
                scenario_counts[key] = scenario_counts.get(key, 0) + 1
    return {
        "scenarios": [k for k, _ in sorted(scenario_counts.items(), key=lambda x: -x[1])[:6]],
        "pros": [k for k, _ in sorted(pro_counts.items(), key=lambda x: -x[1])[:6]],
        "pains": [{"keyword": k, "count": v} for k, v in sorted(pain_counts.items(), key=lambda x: -x[1])[:6]],
        "pain_tags": [k for k, _ in sorted(pain_counts.items(), key=lambda x: -x[1])[:5]],
        "best_review_index": best_review_index,
        "_fallback": True,
    }


def parse_ai_json(raw: str) -> dict:
    text = (raw or "").strip()
    if not text:
        raise ValueError("AI returned empty content")
    think_end = re.search(r'</think\s*>', text, re.IGNORECASE | re.DOTALL)
    if think_end:
        text = text[think_end.end():].strip()
    fence = re.search(r'```(?:json)?\s*\n?(.*?)\n?```', text, re.DOTALL)
    if fence:
        text = fence.group(1).strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        start = text.find("{")
        end = text.rfind("}")
        if start >= 0 and end > start:
            return json.loads(text[start:end + 1])
        raise


def analyze_product_reviews(reviews: list[dict], product_title: str) -> dict:
    """用 MiniMax 对单个产品的评论做语义分析，提取场景/优点/痛点/代表评论"""
    if not reviews:
        return {"scenarios": [], "pros": [], "pains": [], "pain_tags": [], "best_review_index": 0}

    # 构建评论列表文本
    review_lines = []
    for j, r in enumerate(reviews):
        body = (r.get("body_zh", "") or r.get("body", "")).replace("\n", " ").strip()
        title = (r.get("title_zh", "") or r.get("title", "")).strip()
        rating = r.get("rating", "?")
        review_lines.append(f"[{j}] ★{rating} | {title}\n    {body[:300]}")

    prompt = (
        "You are an Amazon product analyst. Analyze the following reviews for the product "
        f"\"{product_title[:80]}\". All reviews are already in Chinese.\n\n"
        "Extract and return ONLY valid JSON (no markdown fences, no extra text):\n"
        "{\n"
        '  "scenarios": ["使用场景1", "使用场景2", ...],\n'
        '  "pros": ["优点1", "优点2", ...],\n'
        '  "pains": [{"keyword": "痛点关键词", "count": N}, ...],\n'
        '  "pain_tags": ["痛点标签1", "痛点标签2", ...],\n'
        '  "best_review_index": N\n'
        "}\n\n"
        "Rules:\n"
        "- scenarios: What user scenarios can be inferred? (e.g. 送礼, 办公, 旅行, 居家, 运动)\n"
        "- pros: What do users consistently praise? Summarize in Chinese, 2-6 items\n"
        "- pains: Pain points with mention count. Merge similar complaints (e.g. 'size too small' + 'runs small' → 尺寸偏小). keyword in Chinese, count as integer\n"
        "- pain_tags: Top 3-5 pain point labels in Chinese\n"
        "- best_review_index: Index (0-based) of the most informative/representative review\n\n"
        "Reviews:\n" + "\n".join(review_lines)
    )

    headers = {"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"}
    payload = {
        "model": "MiniMax-M2.7",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3, "max_tokens": 2000,
    }

    default = fallback_review_insights(reviews)
    last_error = ""
    for attempt in range(AI_MAX_RETRIES):
        try:
            resp = requests.post(
                f"{MINIMAX_BASE_URL}/chat/completions",
                headers=headers, json=payload, timeout=AI_REQUEST_TIMEOUT, verify=False)
            if resp.status_code == 200:
                raw = resp.json()["choices"][0]["message"]["content"].strip()
                result = parse_ai_json(raw)
                # 校验字段
                for k in default:
                    if k not in result:
                        result[k] = default[k]
                return result
            last_error = f"HTTP {resp.status_code}: {resp.text[:200]}"
        except Exception as e:
            last_error = str(e)
            if attempt < AI_MAX_RETRIES - 1:
                time.sleep(min(10, 2 ** attempt * 3))
    default["_ai_error"] = last_error or "AI analysis failed"
    return default


def analyze_cross_product(products: list[dict], category_name: str) -> dict:
    """跨产品分析：合并共性问题 + 生成选品建议"""
    if len(products) < 2:
        # 单产品：直接从其痛点生成简单建议
        suggestions = []
        if products and products[0].get("pains"):
            top_pains = sorted(products[0]["pains"], key=lambda x: -x["count"])[:3]
            top_names = [p["keyword"] for p in top_pains]
            suggestions.append(f'核心改进方向：{"、".join(top_names)}')
        prices = [v for v in (parse_price_value(p.get("price", "")) for p in products) if v is not None]
        if prices:
            suggestions.append(f'价格带参考：${min(prices):.2f} - ${max(prices):.2f}，均价${sum(prices)/len(prices):.2f}')
        return {"common_pains": [], "suggestions": suggestions}

    # 构建跨产品分析 prompt
    products_json = []
    for p in products:
        products_json.append({
            "asin": p.get("asin", ""),
            "title": (p.get("title", "") or "")[:60],
            "price": p.get("price", ""),
            "rating": p.get("rating", ""),
            "review_count": p.get("review_count", 0),
            "monthly_sales": p.get("monthly_sales", ""),
            "pains": p.get("pains", []),
        })
    # 去重保留顺序，避免 prompt 过大
    seen = set()
    unique_products = []
    for p in products_json:
        if p["asin"] not in seen:
            seen.add(p["asin"])
            unique_products.append(p)

    prompt = (
        f"You are an Amazon product analyst. Analyze pain points across {len(unique_products)} products "
        f"in the \"{category_name}\" category. Identify common pain themes and generate actionable sourcing advice.\n\n"
        "Input (JSON array of product analyses):\n"
        + json.dumps(unique_products, ensure_ascii=False, indent=2) + "\n\n"
        "Return ONLY valid JSON (no fences):\n"
        "{\n"
        '  "common_pains": [{"keyword": "共性问题", "total_count": N, "product_count": N}, ...],\n'
        '  "suggestions": ["建议1", "建议2", ...]\n'
        "}\n\n"
        "Rules:\n"
        "- common_pains: Merge similar pains across products, count total mentions and how many products affected. Sort by total_count desc. Max 5 items.\n"
        "- suggestions: Actionable sourcing advice in Chinese, 5-8 items. Each item should be a complete sentence with reasoning, not just a label. Include: price band analysis, core improvement directions (based on top pain points), what users consistently love, the biggest complaint themes to address first, recommended pricing strategy, differentiation opportunities vs competitors, and quality control priorities. Write in the style of an experienced product sourcing consultant."
    )

    headers = {"Authorization": f"Bearer {API_KEY}", "Content-Type": "application/json"}
    payload = {
        "model": "MiniMax-M2.7",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3, "max_tokens": 4000,
    }

    pain_summary = {}
    pain_products = {}
    for p in unique_products:
        asin = p.get("asin", "")
        for pain in p.get("pains", []) or []:
            kw = pain.get("keyword", "")
            cnt = int(pain.get("count", 0) or 0)
            if not kw:
                continue
            pain_summary[kw] = pain_summary.get(kw, 0) + cnt
            pain_products.setdefault(kw, set()).add(asin)
    default = {
        "common_pains": [
            {"keyword": kw, "total_count": cnt, "product_count": len(pain_products.get(kw, set()))}
            for kw, cnt in sorted(pain_summary.items(), key=lambda x: -x[1])[:5]
        ],
        "suggestions": [],
    }
    last_error = ""
    for attempt in range(AI_MAX_RETRIES):
        try:
            resp = requests.post(
                f"{MINIMAX_BASE_URL}/chat/completions",
                headers=headers, json=payload, timeout=AI_REQUEST_TIMEOUT, verify=False)
            if resp.status_code == 200:
                raw = resp.json()["choices"][0]["message"]["content"].strip()
                result = parse_ai_json(raw)
                for k in default:
                    if k not in result:
                        result[k] = default[k]
                return result
            last_error = f"HTTP {resp.status_code}: {resp.text[:200]}"
        except Exception as e:
            last_error = str(e)
            if attempt < AI_MAX_RETRIES - 1:
                time.sleep(min(10, 2 ** attempt * 3))
    default["_ai_error"] = last_error or "AI cross-product analysis failed"
    return default


# ─── JS 提取片段 ─────────────────────────────────────────────────────────────
def get_search_js(domain: str) -> str:
    return f"""
(function(maxCount) {{
    var results = [];
    document.querySelectorAll('[data-asin]').forEach(function(el) {{
        var asin = el.getAttribute('data-asin');
        if (!asin || asin.length !== 10 || !/^B[0-9A-Z]{{9}}$/.test(asin)) return;

        var titleEl = el.querySelector('[id*="productTitle"], h2 a span, h2 span');
        var title = titleEl ? titleEl.textContent.trim().replace(/\\s+/g,' ') : '';

        var price = '';
        var priceEl = el.querySelector('.a-price .a-offscreen');
        if (priceEl) {{ var m = priceEl.textContent.match(/\\$[\\d,]+\\.?\\d*/); if (m) price = m[0]; }}
        if (!price) {{
            var w = el.querySelector('.a-price-whole');
            var f = el.querySelector('.a-price-fraction');
            if (w) price = '$' + w.textContent.trim() + (f ? '.' + f.textContent.trim() : '');
        }}

        var rating = '';
        var rEl = el.querySelector('.a-icon-alt');
        if (rEl) {{ var rm = rEl.textContent.match(/^([\\d.]+)\\s+out of\\s+5/); if (rm) rating = rm[1] + ' out of 5'; }}

        var reviews = '';
        var rvEl = el.querySelector('.a-size-base.s-underline-text');
        if (rvEl) reviews = rvEl.textContent.trim();

        var sales_rank = '';
        var srEl = el.querySelector('.a-section.a-spacing-none .a-size-base.a-color-secondary, [data-component-type="s-product-header-below-secondary-results"] .a-size-base.a-color-secondary');
        if (srEl) {{
            var srText = srEl.textContent.trim();
            var srMatch = srText.match(/#?([\d,]+)\s+in\s+/);
            if (srMatch) sales_rank = '#' + srMatch[1];
        }}

        var brand = '';
        var bEl = el.querySelector('.a-color-secondary, .a-size-base-plus.a-color-base');
        if (bEl) brand = bEl.textContent.trim().replace('Brand: ','');

        var image_url = '';
        var imgEl = el.querySelector('img.s-image');
        if (imgEl) {{
            image_url = imgEl.src || imgEl.getAttribute('data-src') || '';
            if (image_url.indexOf('sprite') !== -1 || image_url.indexOf('data:image') !== -1) image_url = '';
        }}

        // 抓取销量标签 "X+ bought in past month"
        var monthly_sales = '';
        var elText = (el.textContent || '');
        var msMatch = elText.match(/([\d,.]+[KkMm])\+?\s*bought\s+in\s+(?:the\s+)?(?:last|past)\s+month/i);
        if (msMatch) monthly_sales = msMatch[1] + '+';
        if (!monthly_sales && el.parentElement) {{
            var siblings = Array.from(el.parentElement.children);
            for (var s = 0; s < siblings.length && !monthly_sales; s++) {{
                var st = siblings[s].textContent || '';
                var sm = st.match(/([\d,.]+[KkMm])\+?\s*bought\s+in\s+(?:the\s+)?(?:last|past)\s+month/i);
                if (sm) monthly_sales = sm[1] + '+';
            }}
        }}

        if (asin && title) {{
            results.push({{asin, title: title.substring(0,200), brand, price, rating, reviews, sales_rank, monthly_sales,
                          product_url: 'https://{domain}/dp/' + asin,
                          product_image: image_url}});
        }}
    }});
    return JSON.stringify(results.slice(0, maxCount));
}})
"""

REVIEW_EXTRACT_JS = """
(function() {
    var results = [];

    // ── 方法1: data-hook 选择器（Amazon 标准结构，命中率最高）──
    var containers = document.querySelectorAll('[data-hook="review"]');
    for (var i = 0; i < containers.length; i++) {
        var c = containers[i];

        var ratingEl = c.querySelector('[data-hook="review-star-rating"] .a-icon-alt, [data-hook="cmps-review-star-rating"] .a-icon-alt');
        var rating = '?';
        if (ratingEl) {
            var rm = ratingEl.textContent.match(/([\\d.]+)\\s+out of\\s+5/);
            if (rm) rating = rm[1] + ' out of 5 stars';
        }

        var titleEl = c.querySelector('[data-hook="review-title"] span:not(.a-icon-alt)');
        if (!titleEl) titleEl = c.querySelector('[data-hook="review-title"]');
        var title = titleEl ? titleEl.textContent.trim().replace(/^[\\d.]+\\s+out of\\s+5\\s+stars\\s*/i, '') : '';

        var bodyEl = c.querySelector('[data-hook="review-body"] span');
        if (!bodyEl) bodyEl = c.querySelector('[data-hook="review-body"]');
        var body = bodyEl ? bodyEl.textContent.trim() : '';

        var authorEl = c.querySelector('.a-profile-name');
        var author = authorEl ? authorEl.textContent.trim() : 'Anonymous';

        var dateEl = c.querySelector('[data-hook="review-date"]');
        var date = dateEl ? dateEl.textContent.replace('Reviewed in the United States on ', '').trim() : '';

        var verified = !!c.querySelector('[data-hook="avp-badge"]');

        if (title || body) {
            results.push({ rating: rating, title: title, body: body, author: author, date: date, verified: verified });
        }
    }

    // ── 方法2: fallback h5 结构（data-hook 未命中时）──
    if (results.length === 0) {
        var allH5 = document.querySelectorAll('h5');
        for (var i = 0; i < allH5.length; i++) {
            var h5 = allH5[i];
            var text = h5.textContent.trim();
            var match = text.match(/^([\\d.]+)\\s+out of\\s+5\\s+stars\\s+(.+)/i);
            if (!match) continue;
            var parent = h5.closest('li') || h5.parentElement;
            var body = '';
            if (parent) {
                var bodySpan = parent.querySelector('span[class*="review-text"], [data-hook="review-body"] span');
                if (bodySpan) {
                    body = bodySpan.textContent.trim();
                } else {
                    var walker = document.createTreeWalker(parent, NodeFilter.SHOW_TEXT, null, false);
                    var node;
                    while (node = walker.nextNode()) {
                        var t = node.textContent.trim();
                        if (t.length < 20) continue;
                        var tag = node.parentElement ? node.parentElement.tagName : '';
                        if (tag === 'H5' || tag === 'H6' || tag === 'SCRIPT' || tag === 'STYLE' || tag === 'BUTTON') continue;
                        if (t.match(/^\\d+\\.?\\d*\\s+out of\\s+5/i)) continue;
                        body = t; break;
                    }
                }
            }
            var authorEl = parent ? parent.querySelector('a[href*="/profile/"]') : null;
            var author = authorEl ? authorEl.textContent.trim() : 'Anonymous';
            results.push({ rating: match[1]+' out of 5 stars', title: match[2], body: body, author: author, date: '', verified: false });
        }
    }

    return JSON.stringify(results);
})
"""


# ─── Step 1+2: 搜索 ──────────────────────────────────────────────────────────
async def step_search_products(keyword: str, sort: str, max_asins: int, domain: str = "www.amazon.com", region: str = "us") -> list[dict]:
    print(f"\n🔍 Step 1+2: 搜索关键词='{keyword}', sort='{sort}', 目标={max_asins}个产品, 地区={domain}")
    print_localized_keyword_hint(keyword, region)

    port = DEFAULT_PORT
    encoded_keyword = keyword.replace(' ', '+')
    urls = [f"https://{domain}/s?k={encoded_keyword}&s={sort}"]
    if sort == "bestsellers":
        urls.append(f"https://{domain}/s?k={encoded_keyword}")
    await warmup_domain(domain)
    raw = []
    page_error = ""
    for attempt_idx, url in enumerate(urls, 1):
        print(f"   访问: {url}")
        ws = open_new_tab(port, url)
        try:
            await asyncio.sleep(random.uniform(8, 12))

            print("   滚动加载中...")
            for _ in range(random.randint(8, 12)):
                await cdp_send(ws, "Runtime.evaluate",
                               {"expression": "window.scrollTo(0, document.body.scrollHeight)"})
                await asyncio.sleep(random.uniform(1.6, 2.8))

            search_js = get_search_js(domain)
            raw_text = await cdp_send(ws, "Runtime.evaluate",
                                       {"expression": f"({search_js})({max_asins * 3})"})
            page_error = await cdp_send(ws, "Runtime.evaluate", {"expression": """
(() => {
  const text = (document.body && document.body.innerText || '').toLowerCase();
  if (text.includes("we're sorry") || text.includes("an error occurred when we tried to process your request")) return "amazon_error";
  if (text.includes("robot check") || text.includes("enter the characters") || text.includes("captcha")) return "captcha";
  return "";
})()
"""})
        finally:
            close_tab(port, ws)

        try:
            raw = json.loads(raw_text) if raw_text else []
        except Exception:
            raw = []
        if page_error:
            print(f"   ⚠️ Amazon 页面异常: {page_error}，退避后重试")
            await asyncio.sleep(random.uniform(45, 75))
            _WARMED_DOMAINS.discard(domain)
            await warmup_domain(domain)
        if raw or attempt_idx == len(urls):
            break
        print("   ⚠️ 排序搜索未提取到商品，改用普通搜索重试...")
        await asyncio.sleep(random.uniform(4, 8))

    print(f"   JS 提取: {len(raw)} 个候选 ASIN")

    # 去重
    uniq, dup_keys = [], set()
    for p in raw:
        key = (p.get("title","")[:80].lower(), p.get("brand","").lower())
        if key[0] and key not in dup_keys:
            dup_keys.add(key)
            uniq.append(p)

    result_list = uniq[:max_asins]
    print(f"   ✅ 去重完成: 共 {len(uniq)} 个唯一产品,取前 {len(result_list)} 个")
    diagnose_search_drift(result_list, keyword, region)
    return result_list


PRODUCT_DETAIL_JS = """
(function() {
    function text(sel) {
        var el = document.querySelector(sel);
        return el ? el.textContent.trim().replace(/\\s+/g, ' ') : '';
    }
    function attr(sel, name) {
        var el = document.querySelector(sel);
        return el ? (el.getAttribute(name) || '') : '';
    }
    var title = text('#productTitle');
    var brand = '';
    var byline = text('#bylineInfo');
    if (byline) {
        brand = byline.replace(/^Visit the\\s+/i, '').replace(/\\s+Store$/i, '').replace(/^Brand:\\s*/i, '').trim();
    }
    if (!brand) brand = text('tr.po-brand td.a-span9 span, #productOverview_feature_div tr.po-brand td.a-span9 span');
    var price = text('.a-price .a-offscreen') || text('#priceblock_ourprice') || text('#priceblock_dealprice') || text('#corePriceDisplay_desktop_feature_div .a-offscreen');
    var rating = text('#acrPopover .a-icon-alt') || text('[data-hook="rating-out-of-text"]');
    var review_count = text('#acrCustomerReviewText');
    var image = attr('#landingImage', 'src') || attr('#imgTagWrapperId img', 'src');
    var bullets = [];
    document.querySelectorAll('#feature-bullets li span.a-list-item').forEach(function(el) {
        var v = el.textContent.trim().replace(/\\s+/g, ' ');
        if (v && !/make sure this fits/i.test(v)) bullets.push(v);
    });
    var sales_rank = '';
    var detailText = document.body ? document.body.innerText : '';
    var rankMatch = detailText.match(/Best Sellers Rank\\s*[:\\n ]+([\\s\\S]{0,300}?)(?:\\n\\s*Date First Available|\\n\\s*Customer Reviews|\\n\\s*ASIN|$)/i);
    if (rankMatch) {
        sales_rank = rankMatch[1].replace(/\\s+/g, ' ').trim();
    }
    var asin = text('#ASIN') || attr('input#ASIN', 'value');
    return JSON.stringify({title, brand, price, rating, review_count, product_image: image, sales_rank, bullet_points: bullets.slice(0, 8), asin});
})()
"""


async def enrich_product_details(products: list[dict], domain: str = "www.amazon.com") -> list[dict]:
    if not products:
        return products
    print(f"\n🔎 Step 2.5: 商品详情页补全 ({len(products)} 个产品)")
    port = DEFAULT_PORT
    enriched = []
    for i, product in enumerate(products, 1):
        asin = product.get("asin", "")
        if not asin:
            enriched.append(product)
            continue
        url = f"https://{domain}/dp/{asin}"
        print(f"  [{i}/{len(products)}] 补全 {asin}")
        ws = open_new_tab(port, url)
        try:
            await asyncio.sleep(random.uniform(5, 8))
            raw_text = await cdp_send(ws, "Runtime.evaluate", {"expression": PRODUCT_DETAIL_JS})
            detail = json.loads(raw_text) if raw_text else {}
        except Exception as e:
            print(f"    ⚠️ 详情页补全失败: {e}")
            detail = {}
        finally:
            close_tab(port, ws)
        merged = dict(product)
        for key in ("title", "price", "rating", "review_count", "sales_rank", "product_image"):
            if detail.get(key):
                merged[key] = detail[key]
        detail_brand = detail.get("brand", "")
        if detail_brand and detail_brand.lower() not in ("sponsored", "ad", "advertisement"):
            merged["brand"] = detail_brand
        elif str(merged.get("brand", "")).lower() == "sponsored":
            merged["brand"] = ""
        if detail.get("bullet_points"):
            merged["bullet_points"] = detail["bullet_points"]
        enriched.append(merged)
        await asyncio.sleep(random.uniform(3, 6))
    return enriched


# ─── Step 3+4: 评论 ──────────────────────────────────────────────────────────
async def step_scrape_reviews(products: list[dict], translate: bool, domain: str = "www.amazon.com", source_lang: str = "English") -> dict:
    print(f"\n📝 Step 3: 评论抓取 (translate={translate}, source_lang={source_lang})")
    results = {}

    # ── 阶段 A：逐个产品抓取评论（Chrome CDP 必须串行）──
    for i, prod in enumerate(products, 1):
        asin = prod["asin"]
        title = prod.get("title","")[:60]
        print(f"\n  [{i}/{len(products)}] {asin} | {title}...")

        port = DEFAULT_PORT
        review_url = (f"https://{domain}/product-reviews/{asin}/"
                      f"ref=cm_cr_arp_d_viewopt_fmt?sortBy=recent&reviewerType=all_reviews"
                      f"&formatType=current_format&pageNumber=1")

        ws = open_new_tab(port, review_url)
        try:
            await asyncio.sleep(random.uniform(6, 9))

            # 滚动加载
            for _ in range(random.randint(12, 16)):
                await cdp_send(ws, "Runtime.evaluate",
                               {"expression": "window.scrollTo(0, document.body.scrollHeight)"})
                await asyncio.sleep(random.uniform(1.4, 2.4))

            # JS DOM 提取
            raw_text = await cdp_send(ws, "Runtime.evaluate",
                                      {"expression": f"({REVIEW_EXTRACT_JS})()"})
        finally:
            close_tab(port, ws)

        try:
            reviews = json.loads(raw_text) if raw_text else []
        except Exception:
            reviews = []

        print(f"    抓取 {len(reviews)} 条评论")

        results[asin] = {
            "asin": asin,
            "title": prod.get("title", ""),
            "brand": prod.get("brand", ""),
            "price": prod.get("price", ""),
            "rating": prod.get("rating", ""),
            "review_count": prod.get("review_count", ""),
            "product_image": prod.get("product_image", ""),
            "reviews": reviews,
            "monthly_sales": prod.get("monthly_sales", ""),
            "sales_rank": prod.get("sales_rank", ""),
            "bullet_points": prod.get("bullet_points", []),
        }
        await asyncio.sleep(random.uniform(4, 8))

    # ── 阶段 B：并行翻译所有产品的评论（限制并发防 529）────────
    if translate:
        print(f"\n🌐 Step 4: 并行翻译中（{len(results)}个产品，并发≤4）...")
        from concurrent.futures import ThreadPoolExecutor, as_completed

        all_tasks = [(asin, info["reviews"], source_lang) for asin, info in results.items() if info["reviews"]]
        with ThreadPoolExecutor(max_workers=4) as executor:
            future_map = {}
            for asin, revs, lang in all_tasks:
                future_map[executor.submit(translate_batch_m2, revs, 5, lang)] = asin

            done_count = 0
            for future in as_completed(future_map):
                asin = future_map[future]
                try:
                    results[asin]["reviews"] = future.result()
                except Exception:
                    pass
                done_count += 1
                print(f"    [{done_count}/{len(future_map)}] {asin} 翻译完成 "
                      f"({len(results[asin]['reviews'])}条)")

    return results


# ─── 产品图片提取 JS ─────────────────────────────────────────────────────────
PRODUCT_IMAGE_JS = """
(function() {
    // 尝试多种选择器获取产品主图
    var selectors = [
        '#landingImage',
        '#imgBlopsFront',
        '.a-dynamic-image',
        '#main-image',
        'img[data-a-dynamic-image]',
        '#rgb img'
    ];
    for (var i = 0; i < selectors.length; i++) {
        var el = document.querySelector(selectors[i]);
        if (el && el.src && el.src.indexOf('sprite') === -1 && el.src.indexOf('pixel') === -1) {
            return el.src;
        }
    }
    // 备选：从 data-a-dynamic-image 属性提取最大图片
    var dyn = document.querySelector('[data-a-dynamic-image]');
    if (dyn) {
        try {
            var imgs = JSON.parse(dyn.getAttribute('data-a-dynamic-image'));
            var keys = Object.keys(imgs);
            if (keys.length > 0) {
                // 返回最大尺寸的那张
                var maxW = 0, maxUrl = '';
                for (var k = 0; k < keys.length; k++) {
                    var w = parseInt(keys[k].split('x')[0]);
                    if (w > maxW) { maxW = w; maxUrl = keys[k]; }
                }
                return maxUrl || imgs[keys[0]][0];
            }
        } catch(e) {}
    }
    return '';
})()
"""

# ─── Step 5: 分析 ─────────────────────────────────────────────────────────────
def step_compute_analytics(all_data: dict, keyword: str = "") -> dict:
    """独立分析层：逐产品 AI 语义分析 + 跨产品分析，结果写回 all_data。
    返回 cross 数据，供后续 step_analyze_docx、step_analyze、save_db 消费。"""
    from concurrent.futures import ThreadPoolExecutor, as_completed
    emoji, cat_name = get_category_display(keyword)

    # Phase 1: 逐产品语义分析
    print(f"\n🧠 AI 语义分析中（{len(all_data)}个产品并行）...")
    futures = {}
    with ThreadPoolExecutor(max_workers=min(len(all_data), 6)) as executor:
        for asin, info in all_data.items():
            revs = info.get("reviews", [])
            if revs:
                futures[executor.submit(analyze_product_reviews, revs, info.get("title", ""))] = asin
            else:
                info["_ai"] = {"scenarios": [], "pros": [], "pains": [], "pain_tags": [], "best_review_index": 0}

        for future in as_completed(futures):
            asin = futures[future]
            try:
                ai = future.result(timeout=AI_REQUEST_TIMEOUT + 30)
            except Exception as e:
                ai = {"_ai_error": str(e), "scenarios": [], "pros": [], "pains": [], "pain_tags": [], "best_review_index": 0}
            all_data[asin]["_ai"] = ai
            icon = "⚠️" if ai.get("_ai_error") else "✅"
            if ai.get("_ai_error"):
                print(f"   ⚠️ {asin}: AI 失败 - {ai.get('_ai_error')[:60]}")
            else:
                print(f"   ✅ {asin}: {len(ai.get('scenarios',[]))}场景 {len(ai.get('pros',[]))}优点 {len(ai.get('pains',[]))}痛点")

    # Phase 2: 跨产品分析
    cross = {"common_pains": [], "suggestions": []}
    if len(all_data) >= 2:
        print(f"🧠 AI 跨产品分析中...")
        products_for_cross = []
        for asin, info in all_data.items():
            ai = info.get("_ai", {})
            products_for_cross.append({
                "asin": asin,
                "title": (info.get("title", "") or "")[:60],
                "price": info.get("price", ""),
                "rating": info.get("rating", ""),
                "review_count": len(info.get("reviews", [])),
                "monthly_sales": info.get("monthly_sales", ""),
                "pains": ai.get("pains", []),
            })
        cross = analyze_cross_product(products_for_cross, cat_name)
        if cross.get("_ai_error"):
            print(f"   ⚠️ 跨产品 AI 分析失败/超时，已降级")
        else:
            print(f"   ✅ 共性痛点: {len(cross.get('common_pains',[]))} 建议: {len(cross.get('suggestions',[]))}")

    # 写回 all_data
    all_data["__cross__"] = cross
    return cross


def step_analyze_docx(all_data: dict, keyword: str = "", sort_name: str = "", domain: str = "www.amazon.com", output_dir: str = "") -> str:
    """生成 Word 报告（纯消费层：从 all_data._ai 和 all_data.__cross__ 读取已有数据）"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    import requests, hashlib, os, threading, io

    emoji, cat_name = get_category_display(keyword)
    cross = all_data.get("__cross__", {})
    IMG_CACHE_DIR = "/tmp/amazon_image_cache"
    os.makedirs(IMG_CACHE_DIR, exist_ok=True)

    img_bytes_cache = {}
    def dl_img_bytes(url):
        if not url: return None
        if url in img_bytes_cache: return img_bytes_cache[url]
        cache_path = os.path.join(IMG_CACHE_DIR, hashlib.md5(url.encode()).hexdigest() + ".jpg")
        if os.path.exists(cache_path) and time.time() - os.path.getmtime(cache_path) < 7*86400:
            with open(cache_path, "rb") as f:
                data = f.read()
            img_bytes_cache[url] = data
            return data
        try:
            r = requests.get(url, headers={"User-Agent":"Mozilla/5.0","Referer":f"https://{domain}/"}, timeout=15)
            if r.status_code == 200:
                img_bytes_cache[url] = r.content
                with open(cache_path, "wb") as f: f.write(r.content)
                return r.content
        except Exception: pass
        return None

    # 多线程预下载所有产品图片
    threads = []
    for url in [info.get("product_image","") for info in all_data.values() if info.get("product_image")]:
        t = threading.Thread(target=lambda u=url: img_bytes_cache.update({u: dl_img_bytes(u) or b""}))
        t.start(); threads.append(t)
    for t in threads: t.join(timeout=20)

    # ── 创建 Word 文档 ────────────────────────────────────────────
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'PingFang SC'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # ── 标题 ──────────────────────────────────────────────────────
    ts = time.strftime("%Y-%m-%d")
    title = doc.add_heading('', level=0)
    run = title.add_run(f'{emoji} Amazon {cat_name} 选品分析报告')
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x2d, 0x34, 0x36)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'关键词：{keyword} | 地区：{domain} | 排行榜：{sort_name} | 产品数：{len(all_data)} | 日期：{ts}')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()

    # ── 逐产品分析 ────────────────────────────────────────────────
    all_reviews_flat = []

    for idx, (asin, info) in enumerate(all_data.items(), 1):
        revs = info.get("reviews", [])
        all_reviews_flat.extend(revs)
        total = len(revs) or 1
        ai = info.get("_ai", {})

        # 评分分布
        dist = {"5":0,"4":0,"3":0,"2":0,"1":0}
        for r in revs:
            m = re.search(r"([1-5])", r.get("rating","0"))
            if m: dist[m.group(1)] += 1
        dist_str = " ".join(f"{s}★{dist[s]/total*100:.0f}%" for s in ["5","4","3","2","1"])

        # AI 分析数据
        scenarios = ai.get("scenarios", [])
        pros = ai.get("pros", [])
        pains = ai.get("pains", [])
        pain_tags = ai.get("pain_tags", [])
        best_idx = ai.get("best_review_index", 0)
        pain_keywords_str = " | ".join(f"{p.get('keyword','')}({p.get('count',0)}条)" for p in pains) if pains else ""

        # ── 产品标题栏 ────────────────────────────────────────────
        h = doc.add_heading(f'📦 #{idx} {asin}', level=1)
        for run in h.runs:
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0x09, 0x84, 0xe3)

        p = doc.add_paragraph()
        run = p.add_run('产品标题：')
        run.font.bold = True; run.font.size = Pt(10.5)
        run = p.add_run(info.get('title','')[:100])
        run.font.size = Pt(10.5)

        p = doc.add_paragraph()
        run = p.add_run('链接：')
        run.font.bold = True; run.font.size = Pt(10.5)
        run = p.add_run(f'https://{domain}/dp/{asin}')
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x09, 0x84, 0xe3)

        # 价格、评分、评论数、月销量
        monthly_sales_raw = info.get("monthly_sales", "")
        if monthly_sales_raw:
            sales_info = f" | 月销量：{monthly_sales_raw} 件"
        else:
            if total >= 100:
                est_sales = total * 10
                sales_info = f" | 估算月销：~{est_sales:,} 件+"
            elif total >= 30:
                sales_info = " | 月销量：数据暂不可用"
            else:
                sales_info = ""
        p = doc.add_paragraph()
        run = p.add_run(f'价格：{info.get("price","N/A")} | 评分：{info.get("rating","N/A")} | 评论数：{total}条{sales_info}')
        run.font.size = Pt(10.5)

        p = doc.add_paragraph()
        run = p.add_run('评分分布：')
        run.font.bold = True; run.font.size = Pt(10.5)
        run = p.add_run(dist_str)
        run.font.size = Pt(10.5); run.font.color.rgb = RGBColor(0x6c, 0x6c, 0x6c)

        if pain_keywords_str:
            p = doc.add_paragraph()
            run = p.add_run('痛点关键词：')
            run.font.bold = True; run.font.size = Pt(10.5)
            run = p.add_run(pain_keywords_str)
            run.font.size = Pt(10.5); run.font.color.rgb = RGBColor(0xe1, 0x70, 0x55)

        # ── 插入产品图片 ──────────────────────────────────────────
        img_url = info.get("product_image","")
        img_bytes = img_bytes_cache.get(img_url)
        if not img_bytes:
            img_bytes = dl_img_bytes(img_url) if img_url else None
        if img_bytes:
            try:
                # python-docx 不支持 WEBP，需转换为 PNG
                from PIL import Image
                pil_img = Image.open(io.BytesIO(img_bytes))
                if pil_img.mode in ("RGBA", "P"):
                    pil_img = pil_img.convert("RGBA")
                png_buf = io.BytesIO()
                pil_img.save(png_buf, format="PNG")
                png_buf.seek(0)
                doc.add_picture(png_buf, width=Inches(1.8))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            except Exception as e:
                print(f"    ⚠ 图片插入失败: {e}")

        # ── 标签行 ────────────────────────────────────────────────
        if scenarios:
            p = doc.add_paragraph()
            run = p.add_run('🎯 用户场景：')
            run.font.bold = True; run.font.size = Pt(10)
            run = p.add_run('、'.join(scenarios))
            run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x09, 0x84, 0xe3)

        if pros:
            p = doc.add_paragraph()
            run = p.add_run('✅ 核心优点：')
            run.font.bold = True; run.font.size = Pt(10)
            run = p.add_run('、'.join(pros))
            run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x00, 0xb8, 0x94)

        if pain_tags:
            p = doc.add_paragraph()
            run = p.add_run('⚠ 痛点标签：')
            run.font.bold = True; run.font.size = Pt(10)
            run = p.add_run('、'.join(pain_tags))
            run.font.size = Pt(10); run.font.color.rgb = RGBColor(0xe1, 0x70, 0x55)
        else:
            p = doc.add_paragraph()
            run = p.add_run('✨ 无明显痛点')
            run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x00, 0xb8, 0x94)

        # ── 典型评论（AI 选出的最有代表性评论）───────────────────
        best = None
        if revs and 0 <= best_idx < len(revs):
            best = revs[best_idx]
        if not best and revs:
            best = max(revs, key=lambda r: len(r.get("body_zh","") or r.get("body","")))
        if best:
            p = doc.add_paragraph()
            run = p.add_run('💬 典型评论')
            run.font.bold = True; run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x00, 0xb8, 0x94)
            p.paragraph_format.space_before = Pt(6)

            best_title = best.get("title_zh","") or best.get("title","")
            best_body = (best.get("body_zh","") or best.get("body","")).replace("\n"," ").strip()
            if len(best_body) > 400: best_body = best_body[:400] + "..."

            p2 = doc.add_paragraph()
            run = p2.add_run(f'⭐ {best.get("rating","?")} | {best_title}')
            run.font.bold = True; run.font.size = Pt(10)

            p2 = doc.add_paragraph()
            run = p2.add_run(best_body)
            run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p2.paragraph_format.left_indent = Cm(0.5)

        # ── 全部评论详情 ──────────────────────────────────────────
        if revs:
            p = doc.add_paragraph()
            run = p.add_run(f'📋 全部评论详情（{total}条）')
            run.font.bold = True; run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(0x6c, 0x5c, 0xe7)
            p.paragraph_format.space_before = Pt(8)

            for j, r in enumerate(revs, 1):
                stars_text = r.get("rating","?")
                m = re.search(r"([1-5])", stars_text)
                star_count = int(m.group(1)) if m else 0
                star_str = "★" * star_count + "☆" * (5 - star_count)
                rev_title = r.get("title_zh","") or r.get("title","")
                rev_body = (r.get("body_zh","") or r.get("body","")).replace("\n"," ").strip()
                if len(rev_body) > 300: rev_body = rev_body[:300] + "..."

                p = doc.add_paragraph()
                run = p.add_run(f'[{j}] ')
                run.font.size = Pt(9.5); run.font.color.rgb = RGBColor(0xb2, 0xbe, 0xc3)
                run = p.add_run(f'{star_str} {rev_title}')
                run.font.size = Pt(10); run.font.bold = True
                if r.get("verified"):
                    run = p.add_run('  [Verified]')
                    run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x00, 0xb8, 0x94)

                p2 = doc.add_paragraph()
                run = p2.add_run(rev_body)
                run.font.size = Pt(9.5); run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                p2.paragraph_format.left_indent = Cm(0.5)
                p2.paragraph_format.space_after = Pt(4)

        # 分隔线
        doc.add_paragraph('─' * 60)

    # ── 共性痛点分析 ──────────────────────────────────────────────
    h = doc.add_heading('🔍 共性痛点分析', level=1)
    for run in h.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0xe1, 0x70, 0x55)

    p = doc.add_paragraph()
    run = p.add_run(f'基于 {len(all_reviews_flat)} 条评论的 AI 语义分析')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    common_pains = cross.get("common_pains", [])
    if common_pains:
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = '痛点'
        hdr[1].text = '频次'
        hdr[2].text = '分布'
        for cell in hdr:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True; run.font.size = Pt(10)

        for cp in common_pains:
            kw = cp.get("keyword", "?")
            tc = cp.get("total_count", 0)
            pc = cp.get("product_count", 0)
            row = table.add_row().cells
            row[0].text = f'{kw}（{pc}个产品）'
            row[1].text = f'{tc}条'
            bar_len = min(tc, 30)
            row[2].text = '█' * bar_len
            for cell in row:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9.5)
    else:
        p = doc.add_paragraph('暂无明显共性痛点')
        p.runs[0].font.color.rgb = RGBColor(0xb2, 0xbe, 0xc3)

    # ── 竞品对比分析 ──────────────────────────────────────────────
    if len(all_data) >= 2:
        doc.add_paragraph()
        h = doc.add_heading('🔄 竞品对比分析', level=1)
        for run in h.runs:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x6c, 0x5c, 0xe7)

        table = doc.add_table(rows=1, cols=6)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        headers = ['产品', '价格', '评分', '评论数', '月销量', 'Top痛点']
        for i, text in enumerate(headers):
            hdr[i].text = text
            for p in hdr[i].paragraphs:
                for run in p.runs:
                    run.font.bold = True; run.font.size = Pt(10)

        for asin, info in all_data.items():
            revs = info.get("reviews", [])
            title_short = info.get("title","")[:50]
            price = info.get("price","N/A")
            rating = info.get("rating","N/A")
            total = len(revs)
            ai = info.get("_ai", {})

            ms_raw = info.get("monthly_sales", "")
            if ms_raw:
                sales_cell = f'{ms_raw} 件'
            else:
                if total >= 100:
                    sales_cell = f'~{total * 10:,} 件+'
                elif total >= 30:
                    sales_cell = '数据暂不可用'
                else:
                    sales_cell = '-'

            # 该产品的 Top 痛点（来自 AI 分析）
            top_pains = sorted(ai.get("pains", []), key=lambda x: -x.get("count", 0))[:5]
            pain_str = "、".join(f"{p.get('keyword','')}({p.get('count',0)})" for p in top_pains) if top_pains else "无明显痛点"

            row = table.add_row().cells
            row[0].text = f'{title_short}\n{asin}'
            row[1].text = price
            row[2].text = f'⭐ {rating}'
            row[3].text = str(total)
            row[4].text = sales_cell
            row[5].text = pain_str
            for cell in row:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

    # ── 选品建议（AI 生成，与文字报告风格一致）───────────────
    doc.add_paragraph()
    h = doc.add_heading('💡 选品建议', level=1)
    for run in h.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x00, 0xb8, 0x94)

    # 优先使用跨产品 AI 分析的建议
    suggestions = list(cross.get("suggestions", []))

    # 价格带补全
    prices = []
    for info in all_data.values():
        v = parse_price_value(info.get("price", ""))
        if v is not None:
            prices.append(v)
    if prices and not any("价格" in s for s in suggestions):
        suggestions.insert(0, f'价格带参考：${min(prices):.2f} - ${max(prices):.2f}，均价${sum(prices)/len(prices):.2f}')

    # AI 建议为空时，用数据兜底
    if not suggestions:
        if prices:
            suggestions.append(f'价格带参考：${min(prices):.2f} - ${max(prices):.2f}，均价${sum(prices)/len(prices):.2f}')
        all_pains = {}
        for info in all_data.values():
            for pain in info.get("_ai", {}).get("pains", []):
                kw = pain.get("keyword", "")
                if kw:
                    all_pains[kw] = all_pains.get(kw, 0) + int(pain.get("count", 0) or 0)
        if all_pains:
            top = sorted(all_pains.items(), key=lambda x: -x[1])[:5]
            suggestions.append(f'核心改进方向：{"、".join(kw for kw, _ in top)}')
        all_pros = set()
        for info in all_data.values():
            for pro in info.get("_ai", {}).get("pros", []):
                if pro.strip():
                    all_pros.add(pro.strip())
        if all_pros:
            suggestions.append(f'用户最爱：{"、".join(list(all_pros)[:5])}')

    for s in suggestions:
        p = doc.add_paragraph(s, style='List Number')
        for run in p.runs:
            run.font.size = Pt(10.5)

    # ── 保存 ──────────────────────────────────────────────────────
    out_dir = Path(output_dir).expanduser() if output_dir else Path("/tmp")
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = str(out_dir / f"report_{time.strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(out_path)
    print(f"\n📄 Word 报告已生成：{out_path}")
    return out_path


def _guess_sender_from_log(account: str = "competitor") -> str:
    """从 OpenClaw 日志中提取最近一条消息的 senderOpenId，用于自动定位提问人。"""
    import glob, re
    log_pattern = "/tmp/openclaw/openclaw-*.log"
    log_files = sorted(glob.glob(log_pattern), reverse=True)
    for log_path in log_files:
        try:
            with open(log_path, "r", encoding="utf-8") as f:
                # 从文件末尾向前搜索，找到最近一条 dispatch 记录
                lines = []
                chunk_size = 10000
                f.seek(0, 2)  # SEEK_END
                pos = f.tell()
                while pos > 0 and len(lines) < 200:
                    read_size = min(chunk_size, pos)
                    pos -= read_size
                    f.seek(pos)
                    chunk = f.read(read_size)
                    lines = chunk.split("\n") + lines
                # 倒序查找
                for line in reversed(lines):
                    if account not in line or "dispatch" not in line:
                        continue
                    m = re.search(r"senderOpenId[:\s]*'([^']+)'", line)
                    if m:
                        return m.group(1)
                    m = re.search(r'"senderOpenId"[:\s]*"([^"]+)"', line)
                    if m:
                        return m.group(1)
        except Exception:
            continue
    return ""


async def send_report_to_feishu(docx_path: str, keyword: str = "", feishu_target: str = "", feishu_account: str = ""):
    """将 Word 报告发送到飞书。返回 "sent" / "failed" """
    import asyncio as aio, shutil

    # 确定目标 open_id：参数 > 环境变量 FEISHU_TARGET > 日志自动提取 > 兜底值
    target = feishu_target or os.environ.get("FEISHU_TARGET", "")
    if not target:
        target = _guess_sender_from_log(feishu_account or "competitor")
        if target:
            print(f"  🔍 从 OpenClaw 日志自动提取 senderOpenId: {target}")
    if not target:
        target = "ou_2e20fbe54f0f207861644fe56396ef78"
        print(f"  ⚠️ 无法获取 senderOpenId，使用兜底值")
    # 确定飞书账号名：参数 > 环境变量 FEISHU_ACCOUNT > 兜底值
    account = feishu_account or os.environ.get("FEISHU_ACCOUNT", "competitor")

    # openclaw message send --media 要求文件在 mediaLocalRoots 允许的目录下
    # 允许: /tmp/openclaw, ~/.openclaw/media, ~/.openclaw/workspace, ~/.openclaw/sandboxes
    media_dir = os.path.expanduser("~/.openclaw/media")
    os.makedirs(media_dir, exist_ok=True)
    fname = os.path.basename(docx_path)
    media_path = os.path.join(media_dir, fname)
    shutil.copy2(docx_path, media_path)

    try:
        cmd = [
            os.path.expanduser("~/.npm-global/bin/openclaw"), "message", "send",
            "--channel", "feishu",
            "--account", account,
            "--message", f"📊 Amazon 选品分析报告 - {keyword}\n📄 Word 完整图文报告（手机可下载打开）",
            "--media", media_path,
        ]
        if target:
            # openclaw message send --channel feishu 要求 target 格式为 user:<open_id>
            if not target.startswith("user:") and not target.startswith("chat:"):
                target = f"user:{target}"
            cmd.extend(["--target", target])

        proc = await aio.create_subprocess_exec(
            *cmd,
            stdout=aio.subprocess.PIPE,
            stderr=aio.subprocess.PIPE,
        )
        stdout, stderr = await aio.wait_for(proc.communicate(), timeout=180)
        if proc.returncode == 0:
            print(f"  ✅ Word 报告已发送到飞书（手机可直接下载）")
            return "sent"
        err = (stderr or stdout or b"").decode(errors="ignore").strip()[:800]
        print(f"  ⚠️ 飞书发送失败(returncode={proc.returncode}): {err}")
    except Exception as e:
        print(f"  ⚠️ 飞书发送失败: {e}")
    finally:
        # 清理复制的临时文件
        if os.path.exists(media_path):
            try:
                os.unlink(media_path)
            except OSError:
                pass

    return "failed"


def step_analyze(all_data: dict, keyword: str = "", brief: bool = False, domain: str = "www.amazon.com") -> str:
    """生成文本分析报告。brief=True 时只输出产品概览+痛点，不输出评论原文"""
    print("\n📊 Step 5: 分析报告中...")
    W = 72
    emoji, cat_name = get_category_display(keyword)

    rpt = ["=" * W, f"  {emoji} Amazon {cat_name} 选品分析报告", "=" * W]

    all_reviews = []

    for idx, (asin, info) in enumerate(all_data.items(), 1):
        revs = info.get("reviews", [])
        all_reviews.extend(revs)
        dist = {"5":0,"4":0,"3":0,"2":0,"1":0}
        for r in revs:
            m = re.search(r"([1-5])", r.get("rating","0"))
            if m: dist[m.group(1)] += 1
        total = len(revs) or 1

        rpt.append(f"\n  📦 #{idx} [{asin}]")
        rpt.append(f"     产品标题: {info.get('title','')[:80]}")
        rpt.append(f"     链接: https://{domain}/dp/{asin}")
        ms_raw = info.get("monthly_sales", "")
        if ms_raw:
            est_str = f" | 月销:{ms_raw}件"
        else:
            total_num = total
            if total_num >= 10000: sf = 25
            elif total_num >= 1000: sf = 20
            elif total_num >= 100: sf = 15
            else: sf = 8
            est = total_num // sf if sf else 0
            est_str = f" | 估算月销:~{est:,}件" if est > 0 else ""
        rpt.append(f"     价格: {info.get('price','N/A')} | 评分: {info.get('rating','N/A')} | 评论:{total}条{est_str}")
        rpt.append(f"     评分分布: " + " ".join(f"{s}★{dist[s]/total*100:.0f}%" for s in ["5","4","3","2","1"]))

        # 使用 AI 分析的痛点
        ai = info.get("_ai", {})
        pains = ai.get("pains", [])
        if pains:
            mentioned = [f"{p.get('keyword','')}:{p.get('count',0)}" for p in pains]
            rpt.append(f"     痛点: " + " | ".join(mentioned))

        if not brief and revs:
            rpt.append("     ── 评论摘要 ──")
            for j, r in enumerate(revs, 1):
                body_raw = r.get("body_zh") or r.get("body", "")
                body = body_raw.replace('\n', ' ').strip()
                title = r.get("title_zh") or r.get("title", "")
                rpt.append(f"     [{j}] ⭐{r.get('rating','?')} | {title}")
                rpt.append(f"         {body}")

    # 共性痛点（来自 AI 跨产品分析）
    rpt.append("\n" + "=" * W + "\n  🔍 共性痛点分析\n" + "=" * W)
    # 收集所有产品的 AI 痛点
    all_pains = {}
    for info in all_data.values():
        ai = info.get("_ai", {})
        for p in ai.get("pains", []):
            kw = p.get("keyword", "")
            cnt = p.get("count", 0)
            all_pains[kw] = all_pains.get(kw, 0) + cnt
    if all_pains:
        for pk, cnt in sorted(all_pains.items(), key=lambda x: -x[1]):
            bar_len = min(cnt, 30)
            rpt.append(f"  {pk:<20s} {'█' * bar_len} ({cnt}条)")
    else:
        rpt.append("  暂无明显共性痛点")

    if brief:
        rpt.append(f"\n  💡 摘要模式 | 共 {len(all_reviews)} 条评论已分析，详情请查看 Word 报告或评论 JSON")

    return "\n".join(rpt)


# ─── 主流程 ──────────────────────────────────────────────────────────────────
async def run(keyword: str, sort: str, sort_name: str, max_products: int, translate: bool, skip_confirm: bool = False, keyword_raw: str = "", domain: str = "www.amazon.com", source_lang: str = "English", region: str = "us", save_db: bool = False, export_obsidian: bool = False, obsidian_vault: str = "", feishu_target: str = "", feishu_account: str = ""):
    print("\n" + "="*60)
    print("  🛒 Amazon 选品评论分析 Pipeline")
    print(f"  🌍 地区: {domain} ({source_lang})")
    print("="*60)

    if not skip_confirm:
        confirm = input(
            f"\n📌 确认以下设置:\n"
            f"   关键词: {keyword}\n"
            f"   地区: {domain}\n"
            f"   排行榜: {sort} ({sort_name})\n"
            f"   产品数: {max_products}\n"
            f"   确认开始吗?(Y/n)\n   > "
        ).strip().lower()
        if confirm in ("n", "no"):
            print("   已取消。")
            return


    ts = time.strftime("%Y%m%d_%H%M%S")
    output_dir = make_run_output_dir(ts)
    products = await step_search_products(keyword, sort, max_products, domain=domain, region=region)
    if not products:
        print("\n❌ 搜索阶段未提取到任何商品，已中止。")
        print("   可能原因：Amazon 临时错误页、验证码/风控、地区弹窗、排序参数不兼容或页面结构变化。")
        print("   建议：在 18800 Chrome 中手动打开搜索页确认页面正常后重试，或换关键词/榜单。")
        return
    products = await enrich_product_details(products[:max_products], domain=domain)
    tmp_asins = str(output_dir / f"asin_pool_{ts}.json")
    with open(tmp_asins, "w") as f:
        json.dump(products, f, ensure_ascii=False, indent=2)

    all_data = await step_scrape_reviews(products[:max_products], translate=translate, domain=domain, source_lang=source_lang)
    tmp_reviews = str(output_dir / f"reviews_{ts}.json")
    with open(tmp_reviews, "w") as f:
        json.dump(all_data, f, ensure_ascii=False, indent=2)

    # 独立分析层：AI 语义分析（结果写回 all_data._ai 和 all_data.__cross__）
    cross_data = step_compute_analytics(all_data, keyword=keyword_raw or keyword)

    # 消费层：生成 Word 报告（从 all_data 读取已有分析结果）
    docx_path = step_analyze_docx(all_data, keyword=keyword_raw or keyword, sort_name=sort_name, domain=domain, output_dir=str(output_dir))

    # 消费层：生成文字报告（从 all_data 读取已有分析结果）
    report_txt = step_analyze(all_data, keyword=keyword_raw or keyword, brief=False, domain=domain)

    db_path = ""
    if save_db:
        try:
            from storage import save_pipeline_run
            db_path = save_pipeline_run(
                all_data,
                keyword=keyword_raw or keyword,
                sort=sort,
                region=region,
                domain=domain,
                docx_path=docx_path,
                run_id=ts,
                cross_data=cross_data,
            )
            print(f"  ✅ 历史快照已保存: {db_path}")
        except Exception as e:
            print(f"  ⚠️ 历史快照保存失败: {e}")

    obsidian_report = ""
    if export_obsidian:
        try:
            from obsidian_export import export_obsidian as export_obsidian_notes
            obsidian_result = export_obsidian_notes(
                all_data=all_data,
                keyword=keyword_raw or keyword,
                region=region,
                sort=sort,
                domain=domain,
                docx_path=docx_path,
                vault_path=obsidian_vault or None,
                run_id=ts,
            )
            obsidian_report = obsidian_result.get("report_path", "")
            print(f"  ✅ Obsidian 笔记已导出: {obsidian_report}")
        except Exception as e:
            print(f"  ⚠️ Obsidian 导出失败: {e}")

    # 发送 Word 报告到飞书
    send_result = await send_report_to_feishu(docx_path, keyword=keyword_raw or keyword, feishu_target=feishu_target, feishu_account=feishu_account)

    # 自动打开 Word 报告
    import subprocess as sp
    try:
        sp.Popen(["open", docx_path])
    except Exception:
        pass

    print(f"\n{'='*72}")
    print(f"  ✅ 流程完成!")
    print(f"  📦 ASIN池: {tmp_asins}")
    print(f"  📝 评论数据: {tmp_reviews}")
    print(f"  📄 Word报告: {docx_path}")
    if db_path:
        print(f"  🗄️ 历史数据库: {db_path}")
    if obsidian_report:
        print(f"  📝 Obsidian笔记: {obsidian_report}")

    # 如果飞书发送被跳过或失败，打印醒目的发送指令供 agent 接管
    if send_result != "sent":
        print(f"\n  ╔══════════════════════════════════════════════════════════════════╗")
        print(f"  ║  📮 AGENT: 请使用以下命令将报告发送给用户:                       ║")
        print(f"  ║  openclaw message send --channel feishu \\                       ║")
        if feishu_target:
            target_display = f"user:{feishu_target}" if not feishu_target.startswith("user:") else feishu_target
            print(f"  ║    --target {target_display} \\".ljust(72) + "║")
        if feishu_account:
            print(f"  ║    --account {feishu_account} \\".ljust(72) + "║")
        print(f"  ║    --media {docx_path} \\".ljust(72) + "║")
        print(f"  ║    --message \"📊 Amazon 选品分析报告 - {keyword_raw or keyword}\"".ljust(72) + "║")
        print(f"  ╚══════════════════════════════════════════════════════════════════╝")

    print(f"{'='*72}\n")
    print(report_txt)


# ─── 交互入口 ────────────────────────────────────────────────────────────────



def interactive_prompt():
    print("")
    print("=" * 60)
    print("  🛒 Amazon 选品评论分析 Pipeline")
    print("=" * 60)

    keyword = input("\n📌 搜索关键词（例如: athletic socks / wireless charger，中文英文均可）:\n   > ").strip()
    while not keyword:
        keyword = input("   ⚠️ 关键词不能为空，请重新输入\n   > ").strip()
    keyword_raw, keyword = auto_translate_keyword(keyword)

    # ── 地区选择 ──────────────────────────────────────────────
    print("")
    print("📌 目标 Amazon 地区：")
    region_list = list(AMAZON_DOMAINS.keys())
    for i, r in enumerate(region_list, 1):
        print(f"   {i:2d}) {r:<6s} {AMAZON_DOMAINS[r]:<22s} {REGION_NAMES[r]}")
    region_choice = input("   > ").strip()
    while not region_choice.isdigit() or int(region_choice) < 1 or int(region_choice) > len(region_list):
        region_choice = input(f"   ⚠️ 请输入 1~{len(region_list)} 的数字并按回车\n   > ").strip()
    region = region_list[int(region_choice) - 1]
    domain = AMAZON_DOMAINS[region]
    source_lang = SOURCE_LANGUAGES[region]

    print("")
    print("📌 排行榜类型（必须选择）：")
    print("   1) bestsellers       销量榜")
    print("   2) newreleases       新品榜")
    print("   3) moversandshakers  飙升榜")
    print("   4) topreview         评论榜")
    sort_map = {"1":"bestsellers","2":"newreleases","3":"moversandshakers","4":"topreview"}
    sort_cn  = {"bestsellers":"销量榜","newreleases":"新品榜","moversandshakers":"飙升榜","topreview":"评论榜"}
    sort_choice = input("   > ").strip()
    while sort_choice not in ("1","2","3","4"):
        sort_choice = input("   ⚠️ 请输入 1~4 的数字并按回车\n   > ").strip()
    sort = sort_map[sort_choice]
    sort_name = sort_cn[sort]

    max_input = input("\n📌 分析产品数量（必须输入数字）：\n   > ").strip()
    while not max_input.isdigit() or int(max_input) < 1:
        max_input = input("   ⚠️ 请输入正整数，例如 3\n   > ").strip()
    max_products = int(max_input)

    print("")
    print("   ✅ 确认：关键词='" + keyword + "' | 地区='" + REGION_NAMES[region] + "' | 排行榜='" + sort_name + "' | 产品数=" + str(max_products))
    return keyword_raw, keyword, sort, sort_name, max_products, domain, source_lang


if __name__ == "__main__":
    if len(sys.argv) == 1:
        keyword_raw, keyword, sort, sort_name, max_products, domain, source_lang = interactive_prompt()
        translate = True
    else:
        parser = argparse.ArgumentParser(description="Amazon Review Pipeline")
        parser.add_argument("--keyword", "-k")
        parser.add_argument("--region", "-r", choices=list(AMAZON_DOMAINS.keys()), default="us",
                            help="目标 Amazon 地区（默认 us）")
        parser.add_argument("--sort", "-s",
                            choices=["bestsellers","newreleases","moversandshakers","topreview"],
                            help="【必填】排行榜类型：bestsellers/newreleases/moversandshakers/topreview")
        parser.add_argument("--max-products", "-n", type=int,
                            help="【必填】分析产品数量（正整数）")
        parser.add_argument("--translate", "-t", action="store_true", default=True)
        parser.add_argument("--no-translate", action="store_true")
        parser.add_argument("--no-save-db", action="store_false", dest="save_db", default=True, help="不保存到本地 SQLite")
        parser.add_argument("--no-export-obsidian", action="store_false", dest="export_obsidian", default=True, help="跳过导出 Obsidian 笔记")
        parser.add_argument("--obsidian-vault", default="", help="Obsidian vault 路径，默认 ~/Documents/Obsidian/Amazon选品")
        parser.add_argument("--feishu-target", default="", help="飞书目标用户 open_id")
        parser.add_argument("--feishu-account", default="", help="OpenClaw 飞书账号名，必须与 feishu-target 所属应用一致")
        parser.add_argument("--yes", "-y", action="store_true", help="Skip confirm prompt")
        args = parser.parse_args()

        # ── 必填校验：keyword / sort / max-products 三个都必须明确传入 ──────────
        missing = []
        if not args.keyword:
            missing.append("--keyword（搜索关键词）")
        if not args.sort:
            missing.append("--sort（排行榜类型：bestsellers/newreleases/moversandshakers/topreview）")
        if not args.max_products:
            missing.append("--max-products（分析产品数量，正整数）")
        if missing:
            print("❌ 以下参数为必填，请先询问用户后再执行：")
            for m in missing:
                print(f"   {m}")
            print("\n💡 正确调用示例：")
            print("   python3 pipeline.py --keyword 'compression socks' --region de --sort bestsellers --max-products 5 --yes")
            sys.exit(1)

        if args.no_translate:
            args.translate = False
        keyword, sort, max_products = args.keyword, args.sort, args.max_products
        keyword_raw, keyword = auto_translate_keyword(keyword)
        skip_confirm = getattr(args, 'yes', False)
        translate = args.translate
        save_db = args.save_db
        export_obsidian = args.export_obsidian
        obsidian_vault = args.obsidian_vault
        feishu_target = args.feishu_target
        feishu_account = args.feishu_account
        sort_name = {"bestsellers":"销量榜","newreleases":"新品榜","moversandshakers":"飙升榜","topreview":"评论榜"}.get(args.sort,"销量榜")
        region = args.region
        domain = AMAZON_DOMAINS[region]
        source_lang = SOURCE_LANGUAGES[region]

    if len(sys.argv) == 1:
        region = next((r for r, d in AMAZON_DOMAINS.items() if d == domain), "us")
        save_db = True  # 默认保存 SQLite
        export_obsidian = True  # 默认导出 Obsidian
        obsidian_vault = ""
        feishu_target = ""
        feishu_account = ""

    asyncio.run(run(keyword, sort, sort_name, max_products, translate, skip_confirm, keyword_raw=keyword_raw, domain=domain, source_lang=source_lang, region=region, save_db=save_db, export_obsidian=export_obsidian, obsidian_vault=obsidian_vault, feishu_target=feishu_target, feishu_account=feishu_account))
